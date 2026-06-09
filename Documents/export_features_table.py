"""
export_features_table.py
Xuất bảng 52 features ra file Word.
Chạy: python Documents/export_features_table.py
Output: Documents/Bang_52_Features.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT_PATH = Path(__file__).parent / "Bang_52_Features.docx"

# ── Dữ liệu ──────────────────────────────────────────────────────────────────

SECTIONS = [
    {
        "title": "1. Dữ liệu điểm cháy (9 features)",
        "subtitle": "Nguồn: FIRMS (Fire Information for Resource Management System) — MODIS/VIIRS",
        "headers": ["Mã feature", "Tên / Ý nghĩa", "Cách thu thập", "Nguồn"],
        "rows": [
            ("fire_lag_1",           "Ô lưới có cháy 1 ngày trước",              "Shift cột nhãn cháy 1 bước theo grid",                         "FIRMS (MODIS/VIIRS)"),
            ("fire_lag_3",           "Ô lưới có cháy 3 ngày trước",              "Shift cột nhãn cháy 3 bước theo grid",                         "FIRMS (MODIS/VIIRS)"),
            ("neighbor_fire_1d",     "Số ô kề có cháy 1 ngày trước",             "Nhân ma trận kề × fire_lag_1",                                 "FIRMS (MODIS/VIIRS)"),
            ("neighbor_fire_3d",     "Số ô kề có cháy 3 ngày trước",             "Nhân ma trận kề × fire_lag_3",                                 "FIRMS (MODIS/VIIRS)"),
            ("neighbor_fire_7d",     "Số ô kề có cháy 7 ngày trước",             "Nhân ma trận kề × fire_lag_7",                                 "FIRMS (MODIS/VIIRS)"),
            ("fire_count_prev_year", "Số vụ cháy tại ô năm trước",               "Đếm fire theo grid, shift 1 năm",                              "FIRMS (MODIS/VIIRS)"),
            ("fire_count_prev_3y",   "Tổng số vụ cháy 3 năm trước",              "Rolling sum 3 năm (shift 1 trước)",                            "FIRMS (MODIS/VIIRS)"),
            ("fire_freq_5y",         "Tần suất cháy trung bình 5 năm",           "Rolling mean 5 năm (shift 1 trước)",                           "FIRMS (MODIS/VIIRS)"),
            ("days_since_last_fire", "Số ngày kể từ lần cháy cuối trong ô",      "Shift + forward-fill trên cột fire",                           "FIRMS (MODIS/VIIRS)"),
        ],
    },
    {
        "title": "2. Dữ liệu địa hình (4 features)",
        "subtitle": "Nguồn: SRTM 30m (NASA) — trích xuất qua Google Earth Engine",
        "headers": ["Mã feature", "Tên / Ý nghĩa", "Cách thu thập", "Nguồn"],
        "rows": [
            ("dem_mean",  "Độ cao trung bình trong ô lưới (m)",      "Tính mean pixel DEM trong vùng ô lưới",         "SRTM 30m (NASA / GEE)"),
            ("dem_stdev", "Độ lệch chuẩn độ cao trong ô lưới",       "Tính stdDev pixel DEM trong vùng ô lưới",       "SRTM 30m (NASA / GEE)"),
            ("slp_mean",  "Độ dốc trung bình trong ô lưới (°)",      "Tính slope từ DEM rồi lấy mean theo ô",         "Dẫn xuất từ SRTM (GEE)"),
            ("slp_stdev", "Độ lệch chuẩn độ dốc trong ô lưới",      "Tính slope từ DEM rồi lấy stdDev theo ô",       "Dẫn xuất từ SRTM (GEE)"),
        ],
    },
    {
        "title": "3. Dữ liệu khí tượng (5 features)",
        "subtitle": "Nguồn: ERA5-Land (ECMWF / Copernicus) — tải qua CDS API",
        "headers": ["Mã feature", "Tên / Ý nghĩa", "Cách thu thập", "Nguồn"],
        "rows": [
            ("tmean", "Nhiệt độ không khí trung bình ngày (°C)",  "Tải biến 2m_temperature từ ERA5-Land",                  "ERA5-Land (ECMWF/Copernicus)"),
            ("rh",    "Độ ẩm tương đối trung bình ngày (%)",       "Tính từ dewpoint + temperature (ERA5)",                 "ERA5-Land (ECMWF/Copernicus)"),
            ("wind",  "Tốc độ gió trung bình ngày (m/s)",          "Tổng hợp u10 + v10 từ ERA5-Land",                      "ERA5-Land (ECMWF/Copernicus)"),
            ("rain",  "Lượng mưa ngày (mm)",                        "Tải biến total_precipitation từ ERA5-Land",            "ERA5-Land (ECMWF/Copernicus)"),
            ("vpd",   "Vapor Pressure Deficit (kPa)",               "Tính từ tmean và rh theo công thức vật lý",            "Dẫn xuất từ ERA5-Land"),
        ],
    },
    {
        "title": "4. Dữ liệu chỉ số thực vật (5 features)",
        "subtitle": "Nguồn: Sentinel-2 Level-2A (ESA / Copernicus) — trích xuất qua Google Earth Engine",
        "headers": ["Mã feature", "Tên / Ý nghĩa", "Cách thu thập", "Nguồn"],
        "rows": [
            ("ndvi",     "Normalized Difference Vegetation Index",  "(NIR − Red) / (NIR + Red), mean theo ô",            "Sentinel-2 (ESA / GEE)"),
            ("ndvi_std", "Độ lệch chuẩn NDVI trong ô lưới",        "stdDev pixel NDVI trong vùng ô lưới",                "Sentinel-2 (ESA / GEE)"),
            ("ndwi",     "Normalized Difference Water Index",       "(Green − NIR) / (Green + NIR), mean theo ô",         "Sentinel-2 (ESA / GEE)"),
            ("nbr",      "Normalized Burn Ratio",                   "(NIR − SWIR2) / (NIR + SWIR2), mean theo ô",         "Sentinel-2 (ESA / GEE)"),
            ("ndii",     "Normalized Difference Infrared Index",    "(NIR − SWIR1) / (NIR + SWIR1), mean theo ô",         "Sentinel-2 (ESA / GEE)"),
        ],
    },
    {
        "title": "5. Dữ liệu liên quan đến con người (10 features)",
        "subtitle": "",
        "headers": ["Mã feature", "Tên / Ý nghĩa", "Cách thu thập", "Nguồn"],
        "rows": [
            ("dist_road_km",        "Khoảng cách đến đường gần nhất (km)",          "Khoảng cách Euclidean từ centroid ô đến geometry đường",      "OpenStreetMap"),
            ("dist_settlement_km",  "Khoảng cách đến khu dân cư (km)",              "Khoảng cách Euclidean từ centroid ô đến geometry dân cư",     "OpenStreetMap"),
            ("dist_forest_edge_km", "Khoảng cách đến rìa rừng (km)",                "Khoảng cách Euclidean đến biên ranh giới rừng",               "OpenStreetMap"),
            ("dist_powerline_km",   "Khoảng cách đến đường dây điện (km)",          "Khoảng cách Euclidean đến geometry đường dây",                "OpenStreetMap"),
            ("lulc_class",          "Loại hình sử dụng đất (mã số)",                "Trích xuất giá trị raster dominant class theo ô",             "ESA WorldCover (GEE)"),
            ("cropland_frac_1km",   "Tỷ lệ đất nông nghiệp trong 1km (0–1)",        "Tính fraction pixel cropland trong bán kính 1km",             "ESA WorldCover (GEE)"),
            ("tree_cover_pct",      "Tỷ lệ che phủ cây năm 2000 (%)",               "Trích xuất giá trị treecover2000 theo ô",                     "Hansen Global Forest Change (GEE)"),
            ("deforestation_lag_1y","Diện tích mất rừng năm trước (ha)",            "Đếm pixel lossyear theo ô và năm",                           "Hansen Global Forest Change (GEE)"),
            ("nightlight_mean",     "Cường độ ánh sáng ban đêm trung bình",         "Tính mean pixel radiance theo ô và năm",                      "VIIRS Nighttime Light — NASA/NOAA (GEE)"),
            ("pop_density",         "Mật độ dân số (người/km²)",                    "Trích xuất giá trị raster theo ô và năm",                     "WorldPop / GPW (GEE)"),
        ],
    },
    {
        "title": "6. Feature Engineering (19 features)",
        "subtitle": "Tính toán nội bộ từ các dữ liệu thô",
        "headers": ["Mã feature", "Tên / Ý nghĩa", "Công thức / Cách tính", "Dữ liệu gốc"],
        "rows": [
            ("rain_14d_sum",         "Tổng lượng mưa 14 ngày trước",                    "Rolling sum 14 ngày theo grid",                                    "rain (ERA5)"),
            ("rain_30d_sum",         "Tổng lượng mưa 30 ngày trước",                    "Rolling sum 30 ngày theo grid",                                    "rain (ERA5)"),
            ("vpd_14d_mean",         "VPD trung bình 14 ngày trước",                    "Rolling mean 14 ngày theo grid",                                   "vpd"),
            ("vpd_30d_mean",         "VPD trung bình 30 ngày trước",                    "Rolling mean 30 ngày theo grid",                                   "vpd"),
            ("sin_doy",              "Thành phần sin của ngày trong năm",               "sin(2π × DOY / 365)",                                              "date"),
            ("cos_doy",              "Thành phần cos của ngày trong năm",               "cos(2π × DOY / 365)",                                              "date"),
            ("neighbor_count",       "Số ô lưới kề cạnh",                               "Đếm từ adjacency matrix (lưới không gian)",                        "Lưới không gian"),
            ("vpd_neighbor_1d",      "Tương tác VPD × lan truyền cháy",                 "vpd × neighbor_fire_1d",                                           "vpd, neighbor_fire_1d"),
            ("vpd_fire_lag_1",       "Tương tác VPD × cháy ngày trước",                 "vpd × fire_lag_1",                                                 "vpd, fire_lag_1"),
            ("delta_ndvi_14d",       "Biến thiên NDVI so với 14 ngày trước",            "NDVI(t−14) − NDVI(t)",                                             "ndvi"),
            ("delta_nbr_7d",         "Biến thiên NBR so với 7 ngày trước",              "NBR(t−7) − NBR(t)",                                                "nbr"),
            ("has_s2",               "Cờ có ảnh Sentinel-2 trong ngày",                 "1 nếu NDVI không null, 0 nếu bị mây/thiếu dữ liệu",               "ndvi"),
            ("burn_season_flag",     "Cờ mùa khô / mùa cháy",                          "1 nếu tháng 1–4, 0 các tháng còn lại",                            "date"),
            ("days_since_harvest",   "Số ngày kể từ kết thúc vụ thu hoạch",            "DOY − 365 (hoặc vòng năm nếu DOY < 365)",                         "date"),
            ("is_pathway_p1",        "Ngữ cảnh đốt nương sau thu hoạch",               "burn_season_flag=1 & days_since_harvest<30",                       "burn_season_flag, days_since_harvest"),
            ("is_pathway_p2",        "Ngữ cảnh phá rừng / khai hoang",                 "deforestation_lag_1y>1.5 & fire_count_prev_year>0",                "deforestation_lag_1y, fire_count_prev_year"),
            ("is_pathway_p3",        "Ngữ cảnh đốt lặp lại theo mùa",                  "fire_count_prev_year>1 & burn_season_flag=1",                      "fire_count_prev_year, burn_season_flag"),
            ("human_pathway_score",  "Số pathway nhân tạo được kích hoạt (0–3)",        "P1 + P2 + P3",                                                     "is_pathway_p1/p2/p3"),
            ("human_signal_strength","Cường độ tín hiệu nhân tạo tổng hợp (0–1)",      "Weighted blend của pathway score",                                 "human_pathway_score"),
        ],
    },
]

# ── Helpers ───────────────────────────────────────────────────────────────────

HEADER_COLOR = RGBColor(0x1F, 0x49, 0x7D)   # dark blue
ALT_ROW_COLOR = "D9E2F3"                     # light blue (hex string for XML)
FONT_NAME = "Times New Roman"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])


def add_section(doc: Document, section: dict):
    # Title
    h = doc.add_heading(section["title"], level=2)
    h.runs[0].font.name = FONT_NAME
    h.runs[0].font.color.rgb = HEADER_COLOR

    # Subtitle
    if section["subtitle"]:
        p = doc.add_paragraph(section["subtitle"])
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.name = FONT_NAME

    headers = section["headers"]
    rows = section["rows"]

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        set_cell_bg(hdr_cells[i], "1F497D")
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            cells[c_idx].text = cell_text
            run = cells[c_idx].paragraphs[0].runs[0]
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            if c_idx == 0:
                run.bold = True
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 1:
                set_cell_bg(cells[c_idx], ALT_ROW_COLOR)

    # Column widths (total ~16cm)
    if len(headers) == 4:
        set_col_widths(table, [3.5, 5.0, 5.0, 3.5])

    doc.add_paragraph()  # spacing


# ── Main ──────────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section_obj = doc.sections[0]
section_obj.left_margin   = Cm(2.5)
section_obj.right_margin  = Cm(2.0)
section_obj.top_margin    = Cm(2.5)
section_obj.bottom_margin = Cm(2.5)

# Document title
title = doc.add_heading("Bảng mô tả 52 Features — Mô hình XGBoost ForestFlameAlert", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.name = FONT_NAME
title.runs[0].font.color.rgb = HEADER_COLOR

intro = doc.add_paragraph(
    "Tổng hợp toàn bộ 52 features sử dụng trong mô hình XGBoost v5, "
    "phân theo 6 nhóm: điểm cháy, địa hình, khí tượng, chỉ số thực vật, "
    "con người và feature engineering."
)
intro.runs[0].font.name = FONT_NAME
intro.runs[0].font.size = Pt(11)
doc.add_paragraph()

for sec in SECTIONS:
    add_section(doc, sec)

doc.save(OUT_PATH)
print(f"Done → {OUT_PATH}")